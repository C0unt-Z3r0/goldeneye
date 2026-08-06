"""
goldeneye/reports/docx_generator.py
Gerador de relatorios DOCX (Word).
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from rich.console import Console

console = Console()


def generate_docx(project_data: dict, output_path: Path) -> Path:
    """Gera relatorio em Word."""
    
    doc = Document()
    
    # Estilos
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # CAPA
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('GOLDENEYE')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xB8, 0x96, 0x0F)
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Relatório de Avaliação de Segurança')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Cliente: {project_data.get("client", "N/A")}\n').bold = True
    info.add_run(f'Alvo: {project_data.get("target", "N/A")}\n')
    info.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n')
    
    doc.add_paragraph()
    
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('⚠️ DOCUMENTO CONFIDENCIAL')
    run.font.color.rgb = RGBColor(0xDC, 0x14, 0x3C)
    run.bold = True
    
    doc.add_page_break()
    
    # SUMARIO EXECUTIVO
    doc.add_heading('1. Sumário Executivo', level=1)
    doc.add_paragraph(project_data.get('executive_summary', 'Relatório gerado pelo Goldeneye.'))
    
    # ATIVOS
    doc.add_heading('2. Ativos Identificados', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = table.rows[0].cells
    hdr[0].text = 'IP'
    hdr[1].text = 'Hostname'
    hdr[2].text = 'Sistema Operacional'
    hdr[3].text = 'Portas Abertas'
    
    for host in project_data.get('hosts', []):
        row = table.add_row()
        row.cells[0].text = host.get('ip', '')
        row.cells[1].text = host.get('hostname', '')
        row.cells[2].text = host.get('os', '')
        row.cells[3].text = host.get('open_ports', '')
    
    doc.add_paragraph()
    
    # VULNERABILIDADES
    doc.add_heading('3. Vulnerabilidades', level=1)
    vulns = project_data.get('vulnerabilities', [])
    if vulns:
        for v in vulns:
            p = doc.add_paragraph()
            run = p.add_run(f"{v.get('severity', 'Info')}: {v.get('name', 'N/A')}")
            run.bold = True
            doc.add_paragraph(f"Descrição: {v.get('description', '')}")
            doc.add_paragraph(f"Impacto: {v.get('impact', '')}")
            doc.add_paragraph(f"Recomendação: {v.get('recommendation', '')}")
            doc.add_paragraph()
    else:
        doc.add_paragraph('Nenhuma vulnerabilidade crítica encontrada.')
    
    # RECOMENDACOES
    doc.add_heading('4. Recomendações', level=1)
    recs = project_data.get('recommendations', [])
    if isinstance(recs, list):
        for r in recs:
            if isinstance(r, dict):
                doc.add_paragraph(f"• {r.get('title', '')}: {r.get('description', '')}")
            else:
                doc.add_paragraph(f"• {r}")
    
    # RODAPE
    doc.add_paragraph()
    doc.add_paragraph(f'Gerado por Goldeneye v1.0 em {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    console.print(f"[green][+] DOCX gerado: {output_path}[/green]")
    return output_path
